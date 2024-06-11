from flask import Flask, send_file, request, jsonify
from flask_cors import CORS
from docx import Document
import os
import io

app = Flask(__name__)
CORS(app)

@app.route('/generate-docx', methods=['POST'])
def generate_docx():
    data = request.json

    if not data:
        return jsonify({"error": "No data provided"}), 400

    replacement_text = data.get('replacement_text', 'first')
    date_text = data.get('date', 'second')
    value1_text = data.get('value1', 'third')
    value2_text = data.get('value2', 'fourth')
    value3_text = data.get('value3', 'fifth')
    value4_text = data.get('value4', 'sixth')
    value5_text = data.get('value5', 'seventh')
    value6_text = data.get('value6', 'eighth')
    value7_text = data.get('value7', 'ninth')
    value8_text = data.get('value8', 'tenth')
    value9_text = data.get('value9', 'eleventh')
    value10_text = data.get('value10', 'twelfth')
    modelSend_text = data.get('modelSend', 'thirteenth')

    base_template_dir = ""
    template_file = ""

    if modelSend_text.lower() == 'plantilla 1':
        base_template_dir = os.path.abspath('templates/model1/')
    elif modelSend_text.lower() == 'plantilla 2':
        base_template_dir = os.path.abspath('templates/model2/')
    elif modelSend_text.lower() == 'plantilla 3':
        base_template_dir = os.path.abspath('templates/model3/')
    elif modelSend_text.lower() == 'plantilla 4':
        base_template_dir = os.path.abspath('templates/model4/')
    elif modelSend_text.lower() == 'plantilla 5':
        base_template_dir = os.path.abspath('templates/model5/')
    elif modelSend_text.lower() == 'plantilla 6':
        base_template_dir = os.path.abspath('templates/model6/')
    else:
        return jsonify({"error": "Invalid modelSend type"}), 400

    if replacement_text.lower() == 'cartas':
        template_file = 'templateCarta.docx'
    elif replacement_text.lower() == 'informes':
        template_file = 'templateInforme.docx'
    elif replacement_text.lower() == 'memorandums':
        template_file = 'templateMemorandum.docx'
    elif replacement_text.lower() == 'circular':
        template_file = 'templateCircular.docx'
    elif replacement_text.lower() == 'instructivos':
        template_file = 'templateInstructivo.docx'
    elif replacement_text.lower() == 'hoja de servicio':
        template_file = 'templateHojaServicio.docx'
    else:
        return jsonify({"error": "Invalid document type"}), 400

    template_path = os.path.join(base_template_dir, template_file)

    if not os.path.exists(template_path):
        return jsonify({"error": "Template file not found"}), 400

    doc = Document(template_path)
    
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, date_text, value1_text, value2_text, value3_text, value4_text, value5_text, value6_text, value7_text, value8_text, value9_text, value10_text, replacement_text.lower())
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, date_text, value1_text, value2_text, value3_text, value4_text, value5_text, value6_text, value7_text, value8_text, value9_text, value10_text, replacement_text.lower())

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return send_file(doc_io, as_attachment=True, download_name='modified_template.docx')

def replace_text_in_paragraph(paragraph, date_text, value1_text, value2_text, value3_text, value4_text, value5_text, value6_text, value7_text, value8_text, value9_text, value10_text, replacement_text):
    if '{date}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{date}', date_text)
    if '{value1}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{value1}', value1_text)
    if '{value2}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{value2}', value2_text)
    if '{value3}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{value3}', value3_text)
    if '{value4}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{value4}', value4_text)
    if '{value5}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{value5}', value5_text)
    if '{value6}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{value6}', value6_text)
    if '{value7}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{value7}', value7_text)
    if replacement_text.lower() == 'hoja de servicio':
        value8_lines = value8_text.split('\n')
        if '{value8_1}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{value8_1}', value8_lines[0])
        if '{value8_2}' in paragraph.text:
            value8_2_text = value8_lines[1] if len(value8_lines) > 1 else ''
            paragraph.text = paragraph.text.replace('{value8_2}', value8_2_text)
        if '{value8_3}' in paragraph.text:
            value8_3_text = value8_lines[2] if len(value8_lines) > 2 else ''
            paragraph.text = paragraph.text.replace('{value8_3}', value8_3_text)
        if '{value8_4}' in paragraph.text:
            value8_4_text = value8_lines[3] if len(value8_lines) > 3 else ''
            paragraph.text = paragraph.text.replace('{value8_4}', value8_4_text)
    elif '{value8}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{value8}', value8_text)
    if replacement_text.lower() == 'hoja de servicio':
        value9_lines = value9_text.split('\n')
        if '{value9_1}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{value9_1}', value9_lines[0])
        if '{value9_2}' in paragraph.text:
            value9_2_text = value9_lines[1] if len(value9_lines) > 1 else ''
            paragraph.text = paragraph.text.replace('{value9_2}', value9_2_text)
        if '{value9_3}' in paragraph.text:
            value9_3_text = value9_lines[2] if len(value9_lines) > 2 else ''
            paragraph.text = paragraph.text.replace('{value9_3}', value9_3_text)
        if '{value9_4}' in paragraph.text:
            value9_4_text = value9_lines[3] if len(value9_lines) > 3 else ''
            paragraph.text = paragraph.text.replace('{value9_4}', value9_4_text)
    if replacement_text.lower() == 'circular':
        value9_lines = value9_text.split('\n')
        if '{value9_1}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{value9_1}', value9_lines[0])
        if '{value9_2}' in paragraph.text:
            value9_2_text = value9_lines[1] if len(value9_lines) > 1 else ''
            paragraph.text = paragraph.text.replace('{value9_2}', value9_2_text)
    elif '{value9}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{value9}', value9_text)
    if replacement_text.lower() == 'circular':
        value10_lines = value10_text.split('\n')
        if '{value10_1}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{value10_1}', value10_lines[0])
        if '{value10_2}' in paragraph.text:
            value10_2_text = value10_lines[1] if len(value10_lines) > 1 else '' 
            paragraph.text = paragraph.text.replace('{value10_2}', value10_2_text)
    elif '{value10}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{value10}', value10_text)
        
if __name__ == '__main__':
    app.run(debug=True)
